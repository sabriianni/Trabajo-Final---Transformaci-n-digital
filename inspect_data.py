import os
import sys

print("Python version:", sys.version)

try:
    import pandas as pd
    print("Pandas is installed.")
except ImportError:
    print("Pandas is NOT installed.")
    pd = None

try:
    import openpyxl
    print("Openpyxl is installed.")
except ImportError:
    print("Openpyxl is NOT installed.")
    openpyxl = None

try:
    import docx
    print("python-docx is installed.")
except ImportError:
    print("python-docx is NOT installed.")
    docx = None

# Let's inspect the Excel file if we can
excel_path = "dataset_alojamientos_transformacion_digital.xlsx"
if os.path.exists(excel_path):
    print(f"\nFound Excel file: {excel_path}")
    if pd is not None:
        try:
            xl = pd.ExcelFile(excel_path)
            print("Sheets:", xl.sheet_names)
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                print(f"\nSheet '{sheet}' shape: {df.shape}")
                print("Columns:", list(df.columns))
                print("First 5 rows:")
                print(df.head())
        except Exception as e:
            print("Error reading with pandas:", e)
    else:
        print("Pandas not available to read excel.")
else:
    print(f"Excel file not found at {excel_path}")

# Let's inspect the docx file if we can
docx_path = "TRABAJO FINAL Transformacion Digital.docx"
if os.path.exists(docx_path):
    print(f"\nFound Word file: {docx_path}")
    if docx is not None:
        try:
            doc = docx.Document(docx_path)
            print(f"Paragraphs: {len(doc.paragraphs)}")
            print("First 20 paragraphs:")
            for i, p in enumerate(doc.paragraphs[:20]):
                if p.text.strip():
                    print(f"- {p.text[:100]}")
        except Exception as e:
            print("Error reading with docx:", e)
    else:
        print("python-docx not available to read word document.")
else:
    print(f"Word file not found at {docx_path}")
